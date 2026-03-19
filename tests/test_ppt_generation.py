"""Tests for PPT generation agent."""
import os
import tempfile
import pytest
from unittest.mock import patch
from backend.agents.ppt_generation import (
    ppt_generation_node, SLIDE_HANDLERS,
    _generate_word_document, _generate_pdf_document,
)
from backend.graph.state import Stage


def _make_sample_slides():
    return [
        {
            "slide_type": "title",
            "topic": "Test Paper",
            "title": "A Novel Approach to Testing",
            "subtitle": "Authors et al. - March 2026",
            "body_points": ["ArXiv: 1234.56789"],
            "speaker_notes": "Welcome to this presentation.",
            "order": 1,
            "image_path": None,
            "image_caption": None,
        },
        {
            "slide_type": "problem",
            "topic": "Test Paper",
            "title": "Problem Statement",
            "subtitle": "Challenges",
            "body_points": [
                "Current methods suffer from high latency in production environments",
                "Existing approaches require large amounts of labeled data",
                "No prior work addresses the combination of efficiency and accuracy",
            ],
            "speaker_notes": "Let's discuss the core problem.",
            "order": 2,
            "image_path": None,
            "image_caption": None,
        },
        {
            "slide_type": "methodology",
            "topic": "Test Paper",
            "title": "Proposed Method",
            "subtitle": "Architecture",
            "body_points": [
                "We propose a transformer-based architecture with novel attention",
                "The model uses multi-scale feature extraction",
                "Training uses a custom loss function combining cross-entropy and contrastive loss",
            ],
            "speaker_notes": "Our approach is novel in several ways.",
            "order": 3,
            "image_path": None,
            "image_caption": None,
        },
        {
            "slide_type": "results",
            "topic": "Test Paper",
            "title": "Experimental Results",
            "subtitle": "Benchmarks",
            "body_points": [
                "Achieves 95.2% accuracy on GLUE benchmark",
                "Outperforms BERT by 3.1 percentage points",
                "Training time reduced by 40% compared to baseline",
            ],
            "speaker_notes": "Results demonstrate clear improvements.",
            "order": 4,
            "image_path": None,
            "image_caption": None,
        },
        {
            "slide_type": "conclusion",
            "topic": "Test Paper",
            "title": "Conclusions & Future Work",
            "subtitle": None,
            "body_points": [
                "Novel method achieves state-of-the-art results",
                "Future work includes multi-language support",
            ],
            "speaker_notes": "Thank you for your attention.",
            "order": 5,
            "image_path": None,
            "image_caption": None,
        },
    ]


_SAMPLE_PAPER = {
    "title": "A Novel Approach to Testing: Transformer-Based Methods for Software Verification",
    "authors": ["Alice Johnson", "Bob Smith", "Charlie Lee"],
    "arxiv_id": "2401.12345",
    "abstract": (
        "We present a novel approach to software testing using transformer-based methods. "
        "Our system achieves state-of-the-art results on multiple benchmarks, including "
        "a 95.2% accuracy on GLUE and a 40% reduction in training time compared to prior work. "
        "The proposed architecture combines multi-scale feature extraction with a custom loss "
        "function that balances cross-entropy and contrastive objectives."
    ),
}


class TestPptGenerationNode:
    def test_generates_pptx(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            state = {
                "session_id": "test-ppt-gen-1",
                "approved_slides": _make_sample_slides(),
                "processed_paper": _SAMPLE_PAPER,
                "slide_contents": [],
            }
            result = ppt_generation_node(state)

            assert result["current_stage"] == Stage.AWAITING_FINAL_REVIEW
            ppt = result["generated_ppt"]
            assert ppt is not None
            assert ppt["slide_count"] == 5
            assert os.path.exists(ppt["file_path"])
            assert ppt["file_path"].endswith(".pptx")
            # Check file is a valid pptx (not empty)
            assert os.path.getsize(ppt["file_path"]) > 5000

    def test_generates_docx(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            state = {
                "session_id": "test-ppt-gen-2",
                "approved_slides": _make_sample_slides(),
                "processed_paper": _SAMPLE_PAPER,
                "slide_contents": [],
            }
            result = ppt_generation_node(state)
            ppt = result["generated_ppt"]
            doc_path = ppt.get("doc_path")
            assert doc_path is not None
            assert os.path.exists(doc_path)
            assert doc_path.endswith(".docx")

    def test_generates_pdf(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            state = {
                "session_id": "test-ppt-gen-pdf",
                "approved_slides": _make_sample_slides(),
                "processed_paper": _SAMPLE_PAPER,
                "slide_contents": [],
            }
            result = ppt_generation_node(state)
            ppt = result["generated_ppt"]
            pdf_path = ppt.get("pdf_path")
            assert pdf_path is not None
            assert os.path.exists(pdf_path)
            assert pdf_path.endswith(".pdf")
            # PDF should have real content
            assert os.path.getsize(pdf_path) > 2000

    def test_all_three_outputs(self, tmp_path):
        """Verify PPTX, DOCX, and PDF are all generated in one run."""
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            state = {
                "session_id": "test-all-outputs",
                "approved_slides": _make_sample_slides(),
                "processed_paper": _SAMPLE_PAPER,
                "slide_contents": [],
            }
            result = ppt_generation_node(state)
            ppt = result["generated_ppt"]
            assert os.path.exists(ppt["file_path"])
            assert os.path.exists(ppt["doc_path"])
            assert os.path.exists(ppt["pdf_path"])

    def test_falls_back_to_slide_contents(self, tmp_path):
        """If no approved_slides, should use slide_contents."""
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            slides = _make_sample_slides()
            state = {
                "session_id": "test-ppt-gen-3",
                "approved_slides": [],
                "slide_contents": slides,
                "processed_paper": {"title": "Test", "authors": [], "arxiv_id": "1234.56789", "abstract": ""},
            }
            result = ppt_generation_node(state)
            assert result["generated_ppt"]["slide_count"] == len(slides)


class TestSlideHandlers:
    def test_all_types_have_handlers(self):
        expected_types = [
            "title", "problem", "background", "contribution", "methodology",
            "architecture", "equation", "algorithm", "experiments", "results",
            "analysis", "discussion", "limitations", "future", "conclusion",
        ]
        for stype in expected_types:
            assert stype in SLIDE_HANDLERS, f"Missing handler for slide type: {stype}"


class TestWordDocument:
    """Tests for professional Word document generation."""

    def test_docx_created_and_valid(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-1", _make_sample_slides(), _SAMPLE_PAPER)
            assert os.path.exists(doc_path)
            assert os.path.getsize(doc_path) > 5000

    def test_docx_contains_title(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-title", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert _SAMPLE_PAPER["title"] in all_text

    def test_docx_contains_authors(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-authors", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Alice Johnson" in all_text

    def test_docx_contains_abstract(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-abstract", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "transformer-based methods" in all_text

    def test_docx_has_toc(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-toc", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Table of Contents" in all_text

    def test_docx_has_section_numbers(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-sections", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            # Should have numbered headings like "1. Problem Statement"
            numbered = [h for h in headings if h and h[0].isdigit()]
            assert len(numbered) >= 3

    def test_docx_has_discussion_sections(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-discussion", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            assert any("Key Points" in h for h in headings)
            assert any("Discussion" in h for h in headings)

    def test_docx_has_styled_fonts(self, tmp_path):
        from docx import Document
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document("test-doc-fonts", _make_sample_slides(), _SAMPLE_PAPER)
            doc = Document(doc_path)
            # Check that Normal style uses Calibri
            normal = doc.styles['Normal']
            assert normal.font.name == 'Calibri'

    def test_docx_handles_empty_paper(self, tmp_path):
        """Should not crash with minimal paper data."""
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            doc_path = _generate_word_document(
                "test-doc-empty",
                _make_sample_slides(),
                {"title": "Minimal", "authors": [], "arxiv_id": "", "abstract": ""},
            )
            assert os.path.exists(doc_path)


class TestPdfDocument:
    """Tests for professional PDF generation."""

    def test_pdf_created_and_valid(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document("test-pdf-1", _make_sample_slides(), _SAMPLE_PAPER)
            assert os.path.exists(pdf_path)
            assert pdf_path.endswith(".pdf")
            assert os.path.getsize(pdf_path) > 2000

    def test_pdf_is_valid_pdf_format(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document("test-pdf-format", _make_sample_slides(), _SAMPLE_PAPER)
            # Valid PDF starts with %PDF
            with open(pdf_path, "rb") as f:
                header = f.read(5)
            assert header == b"%PDF-"

    def test_pdf_contains_title(self, tmp_path):
        """PDF text should include the paper title."""
        import fitz
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document("test-pdf-title", _make_sample_slides(), _SAMPLE_PAPER)
            doc = fitz.open(pdf_path)
            all_text = ""
            for page in doc:
                all_text += page.get_text()
            doc.close()
            assert "Novel Approach" in all_text

    def test_pdf_contains_sections(self, tmp_path):
        """PDF should have section headings from slides."""
        import fitz
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document("test-pdf-sections", _make_sample_slides(), _SAMPLE_PAPER)
            doc = fitz.open(pdf_path)
            all_text = ""
            for page in doc:
                all_text += page.get_text()
            doc.close()
            assert "Problem Statement" in all_text
            assert "Proposed Method" in all_text
            assert "Experimental Results" in all_text

    def test_pdf_contains_bullet_points(self, tmp_path):
        import fitz
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document("test-pdf-bullets", _make_sample_slides(), _SAMPLE_PAPER)
            doc = fitz.open(pdf_path)
            all_text = ""
            for page in doc:
                all_text += page.get_text()
            doc.close()
            assert "95.2%" in all_text
            assert "GLUE" in all_text

    def test_pdf_has_multiple_pages(self, tmp_path):
        import fitz
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document("test-pdf-pages", _make_sample_slides(), _SAMPLE_PAPER)
            doc = fitz.open(pdf_path)
            # Title page + TOC + content = at least 3 pages
            assert len(doc) >= 3
            doc.close()

    def test_pdf_handles_special_chars(self, tmp_path):
        """PDF should handle <, >, & in text without crashing."""
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            slides = _make_sample_slides()
            slides[1]["body_points"] = [
                "Performance: accuracy > 95% & recall < 90%",
                "Formula: loss = -log(p) where p > 0",
            ]
            pdf_path = _generate_pdf_document("test-pdf-special", slides, _SAMPLE_PAPER)
            assert os.path.exists(pdf_path)
            assert os.path.getsize(pdf_path) > 1000

    def test_pdf_handles_empty_paper(self, tmp_path):
        with patch("backend.agents.ppt_generation.settings") as mock_settings:
            mock_settings.output_dir = str(tmp_path)
            pdf_path = _generate_pdf_document(
                "test-pdf-empty",
                _make_sample_slides(),
                {"title": "Minimal", "authors": [], "arxiv_id": "", "abstract": ""},
            )
            assert os.path.exists(pdf_path)
